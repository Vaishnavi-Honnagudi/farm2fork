"""
Traceability Report Generator — produces a DOCX report for a farm produce batch.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

STATUS_COLORS = {
    "PASS": RGBColor(0x15, 0x7A, 0x47),
    "FAIL": RGBColor(0xC0, 0x39, 0x2B),
    "NOT_IN_DB": RGBColor(0xD3, 0x8B, 0x00),
}

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def _add_kv_row(table, key, value):
    row = table.add_row()
    row.cells[0].text = key
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    row.cells[1].text = str(value)

def generate_report(farm_data: dict, compliance: dict, llm_narrative: str, output_dir: str = "reports") -> str:
    """
    Generate a DOCX traceability report.
    Returns path to the generated file.
    """
    os.makedirs(output_dir, exist_ok=True)

    batch_id = farm_data.get("batch_id", "BATCH001")
    crop = farm_data.get("crop", "Unknown")
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"TraceReport_{crop.replace(' ','_')}_{batch_id}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── HEADER BAND ──────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("🌾  Farm-to-Fork Traceability Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"Batch ID: {batch_id}   |   Crop: {crop.title()}   |   Generated: {datetime.date.today()}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ── SECTION 1: FARM DETAILS ───────────────────────────────────
    _add_heading(doc, "1. Farm & Batch Information", level=1, color=RGBColor(0x1A, 0x5C, 0x2E))

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    fields = [
        ("Farmer Name",      farm_data.get("farmer_name", "—")),
        ("Farm Location",    farm_data.get("location", "—")),
        ("State / District", farm_data.get("state", "—")),
        ("Crop",             crop.title()),
        ("Variety",          farm_data.get("variety", "—")),
        ("Area Cultivated",  farm_data.get("area", "—")),
        ("Sowing Date",      farm_data.get("sowing_date", "—")),
        ("Harvest Date",     farm_data.get("harvest_date", "—")),
        ("Batch ID",         batch_id),
        ("Batch Quantity",   farm_data.get("quantity", "—")),
        ("Water Source",     farm_data.get("water_source", "—")),
        ("Irrigation Type",  farm_data.get("irrigation_type", "—")),
    ]
    for k, v in fields:
        _add_kv_row(table, k, v)

    doc.add_paragraph()

    # ── SECTION 2: INPUT LOG ──────────────────────────────────────
    _add_heading(doc, "2. Agricultural Input Log", level=1, color=RGBColor(0x1A, 0x5C, 0x2E))

    # Fertilisers
    _add_heading(doc, "2.1 Fertilisers Applied", level=2)
    fertilisers = farm_data.get("fertilisers", [])
    if fertilisers:
        ft = doc.add_table(rows=1, cols=4)
        ft.style = 'Table Grid'
        hdr = ft.rows[0].cells
        for i, h in enumerate(["Fertiliser", "Quantity", "Application Date", "Method"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], "D5E8D4")
        for f in fertilisers:
            r = ft.add_row().cells
            r[0].text = f.get("name", "")
            r[1].text = f.get("quantity", "")
            r[2].text = f.get("date", "")
            r[3].text = f.get("method", "")
    else:
        doc.add_paragraph("No fertilisers recorded.")

    doc.add_paragraph()

    # Pesticides
    _add_heading(doc, "2.2 Pesticides / Agrochemicals Applied", level=2)
    pesticides = farm_data.get("pesticides", [])
    if pesticides:
        pt = doc.add_table(rows=1, cols=5)
        pt.style = 'Table Grid'
        ph = pt.rows[0].cells
        for i, h in enumerate(["Pesticide", "Dose (mg/kg)", "Application Date", "PHI (days)", "Target Pest"]):
            ph[i].text = h
            ph[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(ph[i], "D5E8D4")
        for p in pesticides:
            r = pt.add_row().cells
            r[0].text = p.get("name", "")
            r[1].text = str(p.get("dose", ""))
            r[2].text = p.get("date", "")
            r[3].text = str(p.get("phi_days", ""))
            r[4].text = p.get("target", "")
    else:
        doc.add_paragraph("No pesticides recorded.")

    doc.add_paragraph()

    # ── SECTION 3: FSSAI COMPLIANCE ──────────────────────────────
    _add_heading(doc, "3. FSSAI MRL Compliance Check", level=1, color=RGBColor(0x1A, 0x5C, 0x2E))

    overall = compliance.get("overall_compliant", True)
    status_p = doc.add_paragraph()
    status_run = status_p.add_run(
        "✅  OVERALL: COMPLIANT" if overall else "❌  OVERALL: NON-COMPLIANT — CORRECTIVE ACTION REQUIRED"
    )
    status_run.bold = True
    status_run.font.size = Pt(12)
    status_run.font.color.rgb = RGBColor(0x15, 0x7A, 0x47) if overall else RGBColor(0xC0, 0x39, 0x2B)

    ref_p = doc.add_paragraph()
    ref_p.add_run(f"Reference: {compliance.get('fssai_ref', '')}").italic = True

    doc.add_paragraph()

    results = compliance.get("results", [])
    if results:
        ct = doc.add_table(rows=1, cols=6)
        ct.style = 'Table Grid'
        ch = ct.rows[0].cells
        for i, h in enumerate(["Pesticide", "Category", "Applied (mg/kg)", "MRL Limit", "Status", "Remarks"]):
            ch[i].text = h
            ch[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(ch[i], "CFE2F3")

        for res in results:
            row = ct.add_row().cells
            row[0].text = res.get("pesticide", "")
            row[1].text = res.get("category", "")
            row[2].text = str(res.get("dose_applied", ""))
            mrl = res.get("mrl_limit", "N/A")
            row[3].text = str(mrl)
            status = res.get("status", "")
            row[4].text = status
            color = {"PASS": "D5F5E3", "FAIL": "FADBD8", "NOT_IN_DB": "FEF9E7"}.get(status, "FFFFFF")
            _set_cell_bg(row[4], color)
            rr = row[4].paragraphs[0].runs
            if rr:
                rr[0].bold = True
                rr[0].font.color.rgb = STATUS_COLORS.get(status, RGBColor(0,0,0))
            row[5].text = res.get("remark", "")

    doc.add_paragraph()

    # ── SECTION 4: AI NARRATIVE ───────────────────────────────────
    _add_heading(doc, "4. Traceability Narrative", level=1, color=RGBColor(0x1A, 0x5C, 0x2E))
    for para_text in llm_narrative.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ── SECTION 5: CERTIFICATIONS ─────────────────────────────────
    _add_heading(doc, "5. Declarations & Certifications", level=1, color=RGBColor(0x1A, 0x5C, 0x2E))
    declarations = [
        "I declare that all information provided is true and accurate to the best of my knowledge.",
        "All pesticides used are registered under the Insecticides Act, 1968.",
        "Pre-Harvest Interval (PHI) has been strictly observed for all chemical inputs.",
        "Water used for irrigation meets quality standards as per IS 11624.",
        "This batch is fit for human consumption subject to compliance of the above checks.",
    ]
    for d in declarations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(d)

    doc.add_paragraph()

    sig_table = doc.add_table(rows=2, cols=2)
    for i, label in enumerate(["Farmer Signature", "Authorized Inspector Signature"]):
        sig_table.rows[0].cells[i].text = label
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        sig_table.rows[1].cells[i].text = "\n\n____________________\nDate: ___________"

    # ── FOOTER ────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generated by Farm-to-Fork Traceability System | FSSAI Reg. Compliant | {datetime.datetime.now().strftime('%d %b %Y %H:%M')}"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(filepath)
    return filepath
